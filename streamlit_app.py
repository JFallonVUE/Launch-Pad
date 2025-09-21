import os, json, math, io, base64
from pathlib import Path
from typing import List, Dict, Any
import streamlit as st

# Optional deps (app still runs without docx/openai by falling back)
try:
    from docx import Document
except Exception:
    Document = None

OPENAI_OK = False
try:
    from openai import OpenAI
    OPENAI_OK = True
except Exception:
    OPENAI_OK = False

# ---------- Config ----------
DATA_DIR = Path("./data")
DATA_DIR.mkdir(exist_ok=True)
EXPORT_DIR = Path("./exports")
EXPORT_DIR.mkdir(exist_ok=True)

CATALOG_DOCX = Path("./VUE Services 2026.docx")
BIASES_DOCX = Path("./Biases.docx")

# ---------- Simple KB parsing ----------
def parse_services_docx(path: Path) -> List[Dict[str, Any]]:
    if not Document or not path.exists(): return []
    doc = Document(str(path))
    items, current = [], {}
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t: continue
        if p.style and p.style.name and p.style.name.lower().startswith("heading"):
            if current: items.append(current)
            current = {"service_id": t.lower().replace(" ","_"), "name": t, "deliverables": [],
                       "constraints": [], "compatible_biases": [], "price_band": "unknown"}
        elif t.startswith("- "):
            current.setdefault("deliverables", []).append(t[2:])
        elif t.lower().startswith("constraints:"):
            current["constraints"] += [x.strip() for x in t.split(":",1)[1].split(";") if x.strip()]
        elif t.lower().startswith("price:"):
            current["price_band"] = t.split(":",1)[1].strip()
        elif t.lower().startswith("biases:"):
            current["compatible_biases"] = [x.strip().lower() for x in t.split(":",1)[1].split(",")]
        else:
            current.setdefault("deliverables", []).append(t)
    if current: items.append(current)
    return items

def parse_biases_docx(path: Path) -> List[Dict[str, Any]]:
    if not Document or not path.exists(): return []
    doc = Document(str(path))
    out, current = [], {}
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t: continue
        if p.style and p.style.name and p.style.name.lower().startswith("heading"):
            if current: out.append(current)
            key = t.split("—")[0].strip().lower().replace(" ","_")
            current = {"key": key, "name": t.split("—")[0].strip(), "definition": "",
                       "copy_patterns": [], "cadence_patterns": [], "compatible_services": []}
        elif t.lower().startswith("definition:"):
            current["definition"] = t.split(":",1)[1].strip()
        elif t.lower().startswith("copy:"):
            current["copy_patterns"] += [x.strip() for x in t.split(":",1)[1].split(";") if x.strip()]
        elif t.lower().startswith("cadence:"):
            current["cadence_patterns"] += [x.strip() for x in t.split(":",1)[1].split(";") if x.strip()]
        elif t.lower().startswith("compatible:"):
            current["compatible_services"] += [x.strip().lower().replace(" ","_") for x in t.split(":",1)[1].split(",")]
        else:
            if not current.get("definition"):
                current["definition"] = t
            else:
                current["copy_patterns"].append(t)
    if current: out.append(current)
    return out

def load_kb() -> Dict[str, Any]:
    # load from data if present, else parse docx, else fallback sample
    cat_json = DATA_DIR / "catalog.json"
    bias_json = DATA_DIR / "biases.json"
    if cat_json.exists() and bias_json.exists():
        return {"services": json.loads(cat_json.read_text()).get("services", []),
                "biases": json.loads(bias_json.read_text()).get("biases", [])}
    services = parse_services_docx(CATALOG_DOCX)
    biases = parse_biases_docx(BIASES_DOCX)
    if not services or not biases:
        # minimal fallback so the app works immediately
        services = [
            {"service_id":"show_stopper","name":"Show Stopper","deliverables":["Hero set"],"constraints":[],"compatible_biases":["anchoring","novelty"],"price_band":"high"},
            {"service_id":"aerials","name":"Aerials","deliverables":["Drone stills"],"constraints":[],"compatible_biases":["authority","social_proof"],"price_band":"medium"},
            {"service_id":"2d_floor_plan","name":"2D Floor Plan","deliverables":["Schematic"],"constraints":[],"compatible_biases":["fluency"],"price_band":"low"},
            {"service_id":"zillow_3d","name":"Zillow 3D","deliverables":["Tour"],"constraints":[],"compatible_biases":["fluency","mere_exposure"],"price_band":"medium"},
            {"service_id":"virtual_staging","name":"Virtual Staging","deliverables":["Staged photos"],"constraints":["vacant_only"],"compatible_biases":["anchoring"],"price_band":"low"},
            {"service_id":"quick_snaps","name":"Quick Snaps","deliverables":["Fast images"],"constraints":[],"compatible_biases":["mere_exposure","loss_aversion"],"price_band":"low"}
        ]
        biases = [
            {"key":"fluency","name":"Fluency / Cognitive Ease","definition":"Reduce cognitive load.",
             "copy_patterns":["Clear, short lines","Chunk specs"],"cadence_patterns":["Morning","Evening"],
             "compatible_services":["2d_floor_plan","zillow_3d"]},
            {"key":"anchoring","name":"Anchoring","definition":"Lead with signature value.",
             "copy_patterns":["Start with the best"],"cadence_patterns":["Lunch"],"compatible_services":["show_stopper"]},
            {"key":"mere_exposure","name":"Mere Exposure","definition":"Repeat to build familiarity.",
             "copy_patterns":["Series posts"],"cadence_patterns":["Evening"],"compatible_services":["quick_snaps","zillow_3d"]}
        ]
    cat_json.write_text(json.dumps({"services":services}, indent=2), encoding="utf-8")
    bias_json.write_text(json.dumps({"biases":biases}, indent=2), encoding="utf-8")
    return {"services":services,"biases":biases}

KB = load_kb()

# ---------- Simple retriever ----------
def _bofe(text: str) -> Dict[str,float]:
    words = [w.lower() for w in text.split()]
    vec={} 
    for w in words: vec[w]=vec.get(w,0)+1
    return vec

def _cos(a:Dict[str,float], b:Dict[str,float])->float:
    dot = sum(a.get(k,0)*b.get(k,0) for k in set(a)|set(b))
    na = math.sqrt(sum(v*v for v in a.values())) or 1.0
    nb = math.sqrt(sum(v*v for v in b.values())) or 1.0
    return dot/(na*nb)

def retrieve_context(answers: dict, k:int=6) -> Dict[str,Any]:
    q = " ".join([f"{k}:{v}" for k,v in answers.items() if isinstance(v,(str,int,float))])
    qv = _bofe(q or "query")
    s = sorted([(s, _cos(qv,_bofe(s["name"]+" "+" ".join(s.get("deliverables",[]))))) for s in KB["services"]],
               key=lambda x:x[1], reverse=True)[:k]
    b = sorted([(b, _cos(qv,_bofe(b["name"]+" "+b.get("definition","")+" "+" ".join(b.get("copy_patterns",[]))))) for b in KB["biases"]],
               key=lambda x:x[1], reverse=True)[:k]
    return {"services":[x[0] for x in s], "biases":[x[0] for x in b]}

# ---------- Signals ----------
def signals(answers: dict) -> dict:
    size = float(answers.get("interiorSizeSqft", 1500))
    condition = answers.get("conditionBand","average")
    tight = bool(answers.get("tightRooms",False))
    natural = answers.get("naturalLight","good")
    timeline = answers.get("timelinePressure", answers.get("priority","medium"))
    def map_condition(c): return {"pristine":0.2,"updated":0.4,"average":0.5,"dated":0.7,"needs_work":0.9}.get(c,0.5)
    def map_natural(n):  return {"excellent":0.2,"good":0.4,"mixed":0.6,"poor":0.8}.get(n,0.5)
    def map_timeline(t): return {"low":0.3,"medium":0.5,"high":0.7,"urgent":0.9,"speed":0.9,"balance":0.6,"maximize_price":0.4}.get(t,0.5)
    complexity = 0.3*map_condition(condition)+0.2*(1-map_natural(natural))+0.2*(1 if tight else 0)+0.3*(size/4000)
    clarity = 0.6*map_condition(condition)+0.4*(1 if tight else 0)
    momentum = map_timeline(timeline)
    brand = 0.2 + (size/5000.0) + (0.2 if answers.get("propertyType")=="Luxury" else 0.0)
    return {
        "complexity": round(min(1.0, complexity),3),
        "clarityNeed": round(min(1.0, clarity),3),
        "momentumPressure": round(min(1.0, momentum),3),
        "brandLift": round(min(1.0, brand),3),
        "locationEfficiency": 0.5,
    }

# ---------- Guardrails ----------
def enforce_guardrails(answers: dict, services: List[Dict[str,Any]]) -> List[Dict[str,Any]]:
    out = []
    vacant = answers.get("occupancy") == "vacant"
    tight = bool(answers.get("tightRooms"))
    remote = answers.get("likelyBuyer") == "remote_buyer"
    for s in services:
        sid = s.get("service_id","")
        if sid == "virtual_staging" and not vacant and not answers.get("explicitVirtualStagingOK", False):
            continue
        out.append(s)
    if tight and all(s.get("service_id")!="2d_floor_plan" for s in out):
        out.insert(0, {"service_id":"2d_floor_plan","name":"2D Floor Plan","rationale":"Adds schematic clarity for tight rooms."})
    if remote and all(s.get("service_id")!="zillow_3d" for s in out):
        out.insert(0, {"service_id":"zillow_3d","name":"Zillow 3D","rationale":"Remote buyers need tour continuity."})
    return out

# ---------- LLM helpers ----------
def llm_decide(answers: dict, sigs: dict, ctx: dict) -> Dict[str,Any]:
    if OPENAI_OK and (os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")):
        key = os.getenv("OPENAI_API_KEY", st.secrets.get("OPENAI_API_KEY"))
        client = OpenAI(api_key=key)
        sys = ("You are a neutral listing production planner. "
               "Compliance: schools/safety factual only; post-production limited to non-material removals + sky/grass. "
               "Return JSON with 3 stacks (High/Medium/Low) and 3 biases (key,name,definition,why,executionBullets 2–3 items).")
        prompt = {"intake_facts":answers,"signals":sigs,"catalog_snippets":{"services":ctx["services"],"biases":ctx["biases"]}}
        resp = client.chat.completions.create(
            model=os.getenv("OPENAI_MODEL","gpt-4.1-mini"),
            response_format={"type":"json_object"},
            temperature=0.3,
            messages=[{"role":"system","content":sys},{"role":"user","content":json.dumps(prompt)}]
        )
        data = json.loads(resp.choices[0].message.content)
    else:
        # Deterministic fallback (still valid)
        data = {
            "stacks":[
                {"tier":"High","services":[
                    {"service_id":"show_stopper","name":"Show Stopper","rationale":"Flagship visuals."},
                    {"service_id":"aerials","name":"Aerials","rationale":"Context & scale."}
                ],"rationale":"Maximum impact"},
                {"tier":"Medium","services":[
                    {"service_id":"zillow_3d","name":"Zillow 3D","rationale":"Tour continuity."},
                    {"service_id":"2d_floor_plan","name":"2D Floor Plan","rationale":"Schematic clarity."}
                ],"rationale":"Core remote-friendly"},
                {"tier":"Low","services":[
                    {"service_id":"2d_floor_plan","name":"2D Floor Plan","rationale":"Clarity."},
                    {"service_id":"quick_snaps","name":"Quick Snaps","rationale":"Speed."}
                ],"rationale":"Lean & fast"}
            ],
            "biases":[
                {"key":"fluency","name":"Fluency","definition":"Ease of processing","why":"Tight/remote contexts",
                 "executionBullets":["Chunk specs","Short headlines"]},
                {"key":"mere_exposure","name":"Mere Exposure","definition":"Familiarity via repetition","why":"Build trust",
                 "executionBullets":["Series posts","Retargeting"]},
                {"key":"anchoring","name":"Anchoring","definition":"Lead with best","why":"Highlight signature",
                 "executionBullets":["Hero-first","Frame comparisons"]}
            ]
        }
    # apply guardrails to each stack
    stacks=[]
    for stk in data["stacks"]:
        pruned = enforce_guardrails(answers, stk["services"])
        stacks.append({"tier":stk["tier"], "services":pruned, "rationale":stk.get("rationale","")})
    return {"stacks":stacks, "biases":data["biases"]}

def llm_copy(intake: dict, chosen_stack: dict, chosen_bias: dict, ctx: dict) -> Dict[str,Any]:
    # Same pattern: real LLM if key, otherwise deterministic text
    if OPENAI_OK and (os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")):
        key = os.getenv("OPENAI_API_KEY", st.secrets.get("OPENAI_API_KEY"))
        client = OpenAI(api_key=key)
        sys = ("You write neutral listing copy using the chosen bias. "
               "Compliance: schools/safety factual only; post-production limits. Return JSON sections exactly.")
        prompt = {"intake":intake,"chosen_stack":chosen_stack,"chosen_bias":chosen_bias,"kb_context":ctx}
        resp = client.chat.completions.create(
            model=os.getenv("OPENAI_MODEL","gpt-4.1-mini"),
            response_format={"type":"json_object"},
            temperature=0.3,
            messages=[{"role":"system","content":sys},{"role":"user","content":json.dumps(prompt)}]
        )
        return json.loads(resp.choices[0].message.content)
    else:
        return {
            "core_listing_print":{
                "MLS description":"Neutral, factual description...",
                "Flyer/Brochure":{"headline":"Crisp Headline","short narrative":"A few neutral lines.","bulleted specs":["3 beds","2 baths","1,800 sqft"]},
                "Feature Sheet":{"Rooms":{"Kitchen":"Quartz, SS","Living":"View window"},"Upgrades/Brands":["Bosch","Kohler"]}
            },
            "digital_social":{
                "Just Listed":["Post A","Post B"],
                "Open House":["Post A","Post B"],
                "Feature Highlights":["Highlight 1","Highlight 2","Highlight 3"],
                "Under Contract/Sold":["Wrap post"],
                "Video scripts":{"walkthrough_60_120s":"Script...","reels_15_60s":"Short script..."},
                "Single-property website copy":"SPW copy...",
                "Ads":["Variant1","Variant2","Variant3","Variant4","Variant5"]
            },
            "direct_outreach":{
                "New Listing email blast":"Email...",
                "Inquiry response templates":{"SMS":"SMS...", "Email":"Email..."},
                "Open House follow-up email":"Follow-up..."
            },
            "cadence":{"Phase I":["Morning 9–11a","Lunch 12–2p","Evening 5–8p"],
                       "Phase II":["Morning 9–11a","Lunch 12–2p","Evening 5–8p"]},
            "ops_checklists":{"Homeowner Prep":["Declutter","Lights on"],"Run of Show":["Arrival","Coverage"],"Gallery Order":["Exteriors first"],"3D/Plan Placement":["Embed in SPW"],"Retouch Notes":["Non-material removals only"]},
            "kpis":["CTR","Video completion","Lead replies"]
        }

# ---------- DOCX export ----------
def make_docx(filename: str, intake: dict, stacks: List[dict], bias_key: str, copy_pack: dict) -> bytes:
    if not Document:
        # simple text fallback zipped as .docx-like; but prefer python-docx installed
        return ("Install python-docx for Word export.\n"+json.dumps(copy_pack, indent=2)).encode("utf-8")
    from docx import Document as D
    doc = D()
    def H(t,l=1): doc.add_heading(t, level=l)
    def P(t): doc.add_paragraph(t)
    def BL(lst): 
        for it in lst: doc.add_paragraph(str(it), style="List Bullet")
    H("Proposal + Listing Lingo Pack", 0)
    P(f"Mode: Deep Dive")
    P(f"Chosen Bias Plan: {bias_key}")
    H("Stacks (High/Medium/Low)",1)
    for s in stacks:
        H(f"{s['tier']}",2)
        P(s.get("rationale",""))
        for it in s["services"]:
            H(it["name"],3); P(it.get("rationale",""))
    H("I. Core Listing & Print",1)
    cp = copy_pack.get("core_listing_print",{})
    for k,v in cp.items():
        H(k,2)
        if isinstance(v,list): BL(v)
        elif isinstance(v,dict):
            for sk,sv in v.items():
                H(sk,3)
                if isinstance(sv,list): BL(sv)
                elif isinstance(sv,dict):
                    for ssk,ssv in sv.items():
                        H(ssk,4); BL(ssv if isinstance(ssv,list) else [ssv])
                else: P(str(sv))
        else: P(str(v))
    H("II. Digital & Social",1)
    ds = copy_pack.get("digital_social",{})
    for k,v in ds.items():
        H(k,2); BL(v if isinstance(v,list) else [v])
    H("III. Direct Outreach",1)
    do = copy_pack.get("direct_outreach",{})
    for k,v in do.items():
        H(k,2); BL(v if isinstance(v,list) else [v])
    H("Phase I/II & Week-1 Cadence",1)
    cad = copy_pack.get("cadence",{"Phase I":["Morning 9–11a","Lunch 12–2p","Evening 5–8p"],"Phase II":["Morning 9–11a","Lunch 12–2p","Evening 5–8p"]})
    for k,v in cad.items(): H(k,2); BL(v)
    H("Operational Checklists",1)
    ops = copy_pack.get("ops_checklists",{})
    for k,v in ops.items(): H(k,2); BL(v if isinstance(v,list) else [v])
    H("KPIs",1); BL(copy_pack.get("kpis",["CTR","Video completion","Lead replies"]))
    H("Disclaimers",1)
    P("Schools/safety factual only (names, distances, links). Post-production limited to non-material removals + sky/grass.")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ---------- UI ----------
st.set_page_config(page_title="LaunchPad AI – VUE Planner", page_icon="✨", layout="wide")

with st.sidebar:
    st.header("Knowledge Base")
    st.caption("Optionally upload your docs to rebuild KB.")
    up_services = st.file_uploader("VUE Services 2026.docx", type=["docx"])
    up_biases = st.file_uploader("Biases.docx", type=["docx"])
    if st.button("Rebuild KB from uploads") and up_services and up_biases:
        CATALOG_DOCX.write_bytes(up_services.read())
        BIASES_DOCX.write_bytes(up_biases.read())
        # rebuild
        DATA_DIR.joinpath("catalog.json").unlink(missing_ok=True)
        DATA_DIR.joinpath("biases.json").unlink(missing_ok=True)
        globals()["KB"] = load_kb()
        st.success("KB rebuilt from uploaded DOCX files.")

st.title("LaunchPad AI Decision Engine (Streamlit)")
st.write("Lighting (15 Qs) → JSON on screen. Deep Dive (40–50 Qs) → JSON + single .docx export.")

tabs = st.tabs(["Lighting","Deep Dive"])
# ----- Lighting -----
with tabs[0]:
    st.subheader("Lighting – 15 Questions")
    a = {}
    col1,col2,col3 = st.columns(3)
    with col1:
        a["propertyType"] = st.selectbox("Property type", ["SFR","Condo","Townhome","Multi","Luxury","Other"])
        a["beds"] = st.number_input("Beds", 0, 20, 3)
        a["baths"] = st.number_input("Baths", 0.0, 20.0, 2.0, 0.5)
        a["interiorSizeSqft"] = st.number_input("Interior size (sqft)", 0, 20000, 1800)
        a["conditionBand"] = st.selectbox("Condition", ["pristine","updated","average","dated","needs_work"])
    with col2:
        a["tightRooms"] = st.checkbox("Tight/small rooms?")
        a["naturalLight"] = st.selectbox("Natural light", ["excellent","good","mixed","poor"])
        a["occupancy"] = st.selectbox("Occupancy", ["vacant","occupied"])
        a["quirkyFlow"] = st.checkbox("Odd/quirky flow?")
        a["signatureFeature"] = st.text_input("Signature feature", "Corner lot")
    with col3:
        a["likelyBuyer"] = st.selectbox("Likely buyer", ["first_time","move_up","downsizer","investor","remote_buyer","luxury"])
        a["locationPerk"] = st.selectbox("Location perk", ["walkable","commute_anchor","schools","parks","views","quiet","other"])
        a["timelinePressure"] = st.selectbox("Timeline pressure", ["low","medium","high","urgent"])
        a["agentOnCamComfort"] = st.selectbox("Agent on-cam comfort", ["low","medium","high"])
        a["showingWindow"] = st.selectbox("Showing window", ["morning","lunch","evening","open"])
    if st.button("Run Lighting"):
        s = signals(a)
        ctx = retrieve_context(a)
        dec = llm_decide(a, s, ctx)
        st.json({"stacks":dec["stacks"], "biases":dec["biases"]})

# ----- Deep Dive -----
with tabs[1]:
    st.subheader("Deep Dive – 40–50 Questions")
    # reuse Lighting basics
    dd = {}
    dd.update(a)  # prefill from Lighting for convenience; users can change
    dd["priority"] = st.selectbox("Goal priority", ["maximize_price","balance","speed"], index=1)
    dd["domTolerance"] = st.selectbox("DOM tolerance", ["short","medium","long"], index=1)
    dd["agentOnCamComfort"] = st.selectbox("Agent on-cam comfort", ["low","medium","high"], index=1)
    dd["distributionChannels"] = st.multiselect("Distribution channels", ["mls","social","email","ads","video","spw"], default=["mls","social","email","video"])
    dd["accessWindows"] = st.multiselect("Access windows (ops)", ["morning","lunch","evening"], default=["morning","lunch","evening"])
    dd["locationFacts"] = st.text_area("Location facts (factual only: names, distances, links)", "").splitlines()
    dd["upgradesBrands"] = [x.strip() for x in st.text_input("Upgrades/brands (comma-separated)", "").split(",") if x.strip()]
    dd["roomsHighlights"] = st.text_area("Rooms highlights (one per line)", "").splitlines()

    chosen_tier = st.selectbox("Select a tier for export", ["High","Medium","Low"])
    chosen_bias_key = st.selectbox("Select a bias plan", [b["key"] for b in KB["biases"]] or ["fluency"])

    if st.button("Run Deep Dive + Export .docx"):
        s = signals(dd)
        ctx = retrieve_context(dd)
        dec = llm_decide(dd, s, ctx)
        st.json({"stacks":dec["stacks"], "biases":dec["biases"]})
        # find chosen stack/bias
        stack = next((x for x in dec["stacks"] if x["tier"].lower()==chosen_tier.lower()), dec["stacks"][0])
        bias = next((b for b in dec["biases"] if b["key"]==chosen_bias_key), dec["biases"][0])
        copy_pack = llm_copy({"answers":dd,"signals":s}, stack, bias, ctx)
        # attach disclaimers
        copy_pack.setdefault("disclaimers",{})
        copy_pack["disclaimers"].setdefault("schools_safety","School/safety references must remain factual only—names, distances, links.")
        copy_pack["disclaimers"].setdefault("post_production","Post-production limited to non-material removals + sky/grass.")
        # build docx
        filename = f"proposal_{chosen_tier}_{chosen_bias_key}.docx"
        blob = make_docx(filename, dd, dec["stacks"], chosen_bias_key, copy_pack)
        st.download_button("Download Proposal + Listing Lingo Pack (.docx)", data=blob, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.success("Generated .docx per spec (one file).")
