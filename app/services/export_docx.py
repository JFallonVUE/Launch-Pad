from docx import Document
from app.config import settings
from app.models import Intake
import os
from typing import Dict, Any

def _h(doc, text, lvl=1): doc.add_heading(text, level=lvl)
def _p(doc, text): doc.add_paragraph(text)
def _bullets(doc, items):
    for it in items: doc.add_paragraph(it, style="List Bullet")

def build_doc(intake: Intake, copy_pack: Dict[str,Any], chosen_tier: str, chosen_bias: str, job_id: str) -> str:
    os.makedirs(settings.EXPORT_DIR, exist_ok=True)
    doc = Document()
    _h(doc, "Proposal + Listing Lingo Pack", 0)
    _p(doc, f"Mode: Deep Dive | Intake ID: {intake.id}")
    _p(doc, f"Chosen Tier: {chosen_tier} | Bias Plan: {chosen_bias}")
    _h(doc, "Chosen Services & Why", 1)
    stack = next(s for s in intake.stacks if s["tier"].lower()==chosen_tier.lower())
    for s in stack["services"]:
        _h(doc, s["name"], 2)
        _p(doc, s.get("rationale",""))

    _h(doc, "I. Core Listing & Print", 1)
    cp = copy_pack.get("core_listing_print", {})
    for k,v in cp.items():
        _h(doc, k, 2)
        if isinstance(v, list): _bullets(doc, v)
        elif isinstance(v, dict):
            for sk,sv in v.items():
                _h(doc, sk, 3)
                if isinstance(sv, list): _bullets(doc, sv)
                elif isinstance(sv, dict):
                    for ssk, ssv in sv.items():
                        _h(doc, ssk, 4)
                        if isinstance(ssv, list): _bullets(doc, ssv)
                        else: _p(doc, str(ssv))
                else: _p(doc, str(sv))
        else: _p(doc, str(v))

    _h(doc, "II. Digital & Social", 1)
    ds = copy_pack.get("digital_social", {})
    for k,v in ds.items():
        _h(doc, k, 2)
        if isinstance(v, list): _bullets(doc, v)
        elif isinstance(v, dict):
            for sk,sv in v.items():
                _h(doc, sk, 3)
                if isinstance(sv, list): _bullets(doc, sv)
                else: _p(doc, str(sv))
        else: _p(doc, str(v))

    _h(doc, "III. Direct Outreach", 1)
    do = copy_pack.get("direct_outreach", {})
    for k,v in do.items():
        _h(doc, k, 2)
        if isinstance(v, list): _bullets(doc, v)
        elif isinstance(v, dict):
            for sk,sv in v.items():
                _h(doc, sk, 3)
                if isinstance(sv, list): _bullets(doc, sv)
                else: _p(doc, str(sv))
        else: _p(doc, str(v))

    _h(doc, "Phase I/II & Week-1 Cadence", 1)
    cadence = copy_pack.get("cadence", {"Phase I": ["Morning 9–11a","Lunch 12–2p","Evening 5–8p"],
                                        "Phase II": ["Morning 9–11a","Lunch 12–2p","Evening 5–8p"]})
    for k,v in cadence.items():
        _h(doc, k, 2); _bullets(doc, v)

    _h(doc, "Operational Checklists", 1)
    ops = copy_pack.get("ops_checklists", {})
    for k,v in ops.items():
        _h(doc, k, 2)
        _bullets(doc, v if isinstance(v,list) else [str(v)])

    _h(doc, "KPIs (Simple)", 1)
    kpis = copy_pack.get("kpis", ["CTR on listing page", "Video completion rate", "Inquiry response time"])
    _bullets(doc, kpis)

    _h(doc, "Disclaimers", 1)
    disc = copy_pack.get("disclaimers", {})
    for k,v in disc.items():
        _h(doc, k.replace("_"," ").title(), 2)
        _p(doc, v)

    outpath = os.path.join(settings.EXPORT_DIR, f"{job_id}.docx")
    doc.save(outpath)
    return outpath
