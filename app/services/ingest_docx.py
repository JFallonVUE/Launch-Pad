from app.config import settings
import os, json
from typing import Dict, Any, List

try:
    from docx import Document
except Exception:
    Document = None

def _parse_services_docx(path: str) -> List[Dict[str, Any]]:
    if Document is None:
        return []
    doc = Document(path)
    items = []
    current = {}
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if p.style and p.style.name and p.style.name.lower().startswith("heading"):
            if current:
                items.append(current)
            current = {"service_id": t.lower().replace(" ","_"), "name": t, "deliverables": [],
                       "constraints": [], "compatible_biases": [], "price_band": "unknown"}
        elif t.startswith("- "):
            current.setdefault("deliverables", []).append(t[2:])
        elif t.lower().startswith("constraints:"):
            current["constraints"] += [x.strip() for x in t.split(":",1)[1].split(";") if x.strip()]
        elif t.lower().startswith("price:"):
            current["price_band"] = t.split(":",1)[1].strip()
        elif t.lower().startswith("biases:"):
            current["compatible_biases"] = [x.strip().lower() for x in t.split(":",1)[1].split(",")]
        else:
            current.setdefault("deliverables", []).append(t)
    if current:
        items.append(current)
    return items

def _parse_biases_docx(path: str) -> List[Dict[str, Any]]:
    if Document is None:
        return []
    doc = Document(path)
    out = []
    current = {}
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if p.style and p.style.name and p.style.name.lower().startswith("heading"):
            if current:
                out.append(current)
            key = t.split("—")[0].strip().lower().replace(" ","_")
            current = {"key": key, "name": t.split("—")[0].strip(), "definition": "",
                       "copy_patterns": [], "cadence_patterns": [], "compatible_services": []}
        elif t.lower().startswith("definition:"):
            current["definition"] = t.split(":",1)[1].strip()
        elif t.lower().startswith("copy:"):
            current["copy_patterns"] += [x.strip() for x in t.split(":",1)[1].split(";") if x.strip()]
        elif t.lower().startswith("cadence:"):
            current["cadence_patterns"] += [x.strip() for x in t.split(":",1)[1].split(";") if x.strip()]
        elif t.lower().startswith("compatible:"):
            current["compatible_services"] += [x.strip().lower().replace(" ","_") for x in t.split(":",1)[1].split(",")]
        else:
            if not current.get("definition"):
                current["definition"] = t
            else:
                current["copy_patterns"].append(t)
    if current:
        out.append(current)
    return out

def build_kb_files():
    os.makedirs(settings.DATA_DIR, exist_ok=True)
    catalog_path = os.path.join(settings.DATA_DIR, "catalog.json")
    biases_path = os.path.join(settings.DATA_DIR, "biases.json")
    services = _parse_services_docx(settings.CATALOG_DOCX_PATH) if os.path.exists(settings.CATALOG_DOCX_PATH) else []
    biases = _parse_biases_docx(settings.BIASES_DOCX_PATH) if os.path.exists(settings.BIASES_DOCX_PATH) else []

    with open(catalog_path, "w", encoding="utf-8") as f:
        json.dump({"services": services}, f, indent=2, ensure_ascii=False)
    with open(biases_path, "w", encoding="utf-8") as f:
        json.dump({"biases": biases}, f, indent=2, ensure_ascii=False)
